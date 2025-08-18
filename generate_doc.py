import os
import shutil
import tempfile
import zipfile
from typing import List

from docx import Document
from langchain_chroma import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain_groq import ChatGroq
from sentence_transformers import SentenceTransformer


# -----------------------------
# File Helpers
# -----------------------------
def unzip_project(zip_path: str, extract_to: str) -> str:
    with zipfile.ZipFile(zip_path, "r") as zip_ref:
        zip_ref.extractall(extract_to)
    return extract_to


def collect_code_files(root_dir: str) -> List[str]:
    code_exts = [
        ".py", ".js", ".jsx", ".ts", ".tsx", ".java", ".cs", ".cpp", ".c", ".go",
        ".rb", ".php", ".swift", ".kt", ".scala", ".rs", ".dart"
    ]
    files = []
    for dirpath, _, filenames in os.walk(root_dir):
        for f in filenames:
            if any(f.endswith(ext) for ext in code_exts):
                files.append(os.path.join(dirpath, f))
    return files


def read_file(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return ""


# -----------------------------
# LLM + Vectorstore Setup
# -----------------------------
def setup_llm():
    return ChatGroq(model="llama3-8b-8192", temperature=0)


def setup_embeddings():
    return HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")


# -----------------------------
# Core Doc Generation
# -----------------------------
def generate_overview(llm, vectorstore) -> str:
    prompt = ChatPromptTemplate.from_template(
        """You are a senior software engineer. Summarize this codebase at a high functional level.
Provide:
1. Purpose of the application
2. Core features
3. Modules/components
4. API endpoints (if any)
5. UI features (if any)
6. Integrations (DB, APIs, auth, etc.)

Base your answer on the indexed code."""
    )
    retriever = vectorstore.as_retriever(search_kwargs={"k": 8})
    chain = prompt | llm
    docs = retriever.get_relevant_documents("Give me a detailed overview of the codebase")
    return chain.invoke({"input": docs}).content


# -----------------------------
# File-by-File Analysis
# -----------------------------
def analyze_single_file(llm, filepath: str, content: str) -> str:
    prompt = ChatPromptTemplate.from_template(
        """You are a senior software engineer documenting a codebase.
Analyze the following file and produce structured notes.

# File
{filename}

# Code
{code}

# Instructions
1. State the purpose of this file.
2. List important functions, classes, or components and their roles.
3. Explain how it connects to other parts of the system.
4. Note any API routes, UI elements, configs, or special patterns.
5. Mention TODOs, risks, or tech debt if visible.

# Answer"""
    )
    trimmed = content[:4000]  # safeguard against huge files
    ans = (prompt | llm).invoke({"filename": filepath, "code": trimmed}).content
    return f"### {os.path.basename(filepath)}\n\n{ans.strip()}\n"


def generate_file_by_file_doc(llm, project_root: str) -> str:
    files = collect_code_files(project_root)
    sections = ["## File-by-File Analysis\n"]
    for fp in files[:50]:  # limit for performance
        content = read_file(fp)
        if not content.strip():
            continue
        sections.append(analyze_single_file(llm, fp, content))
    return "\n".join(sections)


# -----------------------------
# Main Function
# -----------------------------
def generate_functional_doc(zip_path: str, output_dir: str = "output") -> str:
    if os.path.exists(output_dir):
        shutil.rmtree(output_dir)
    os.makedirs(output_dir, exist_ok=True)

    workdir = tempfile.mkdtemp()
    unzip_project(zip_path, os.path.join(workdir, "repo"))

    # Setup LLM + embeddings
    llm = setup_llm()
    embed = setup_embeddings()

    # Collect code
    files = collect_code_files(os.path.join(workdir, "repo"))
    texts = [read_file(f) for f in files if read_file(f).strip()]
    if not texts:
        raise ValueError("No source files found in the uploaded project.")

    # Build vectorstore
    vectorstore = FAISS.from_texts(texts, embed)

    # Generate documentation parts
    md_parts = []

    # 1. Project Overview
    md_parts.append("# Project Functional Documentation\n")
    overview = generate_overview(llm, vectorstore)
    md_parts.append("## Overview\n" + overview + "\n")

    # 2. File-by-File Analysis
    file_analysis_md = generate_file_by_file_doc(llm, os.path.join(workdir, "repo"))
    md_parts.append(file_analysis_md)

    # Save as Markdown
    md_content = "\n".join(md_parts)
    md_path = os.path.join(output_dir, "functional_doc.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)

    # Save as DOCX
    doc = Document()
    doc.add_heading("Project Functional Documentation", 0)
    doc.add_heading("Overview", level=1)
    doc.add_paragraph(overview)

    doc.add_heading("File-by-File Analysis", level=1)
    for fp in collect_code_files(os.path.join(workdir, "repo"))[:50]:
        content = read_file(fp)
        if not content.strip():
            continue
        analysis = analyze_single_file(llm, fp, content)
        doc.add_heading(os.path.basename(fp), level=2)
        for para in analysis.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para)

    doc_path = os.path.join(output_dir, "functional_doc.docx")
    doc.save(doc_path)

    return f"Documentation generated:\n- {md_path}\n- {doc_path}"
